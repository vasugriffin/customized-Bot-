# doc_qa_bot.py
# LangChain for document handling + ZhipuAI official SDK for answering

import os
import pytesseract
from pdf2image import convert_from_path
from PIL import Image
from docx import Document as DocxDocument  # for Word files

from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.schema import Document

from zhipuai import ZhipuAI


# 🔹 Path to Tesseract OCR executable (adjust if needed)
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# 🔹 Set your ZhipuAI API key
os.environ["ZHIPUAI_API_KEY"] = "0b91fae618904fe5bd89cc1b38ff9e6d.DUxRHHzMxULNSpI7"


# OCR for scanned PDFs
def load_scanned_pdf(file_path):
    text = ""
    pages = convert_from_path(file_path, dpi=300)
    for page in pages:
        text += pytesseract.image_to_string(page, lang="eng") + "\n"
    return text


# OCR for images
def load_images(folder_path):
    documents = []
    for file in os.listdir(folder_path):
        if file.lower().endswith((".png", ".jpg", ".jpeg")):
            file_path = os.path.join(folder_path, file)
            try:
                text = pytesseract.image_to_string(Image.open(file_path), lang="eng")
                if text.strip():
                    documents.append(Document(page_content=text, metadata={"source": file}))
                    print(f"🖼️ Extracted text from image: {file}")
            except Exception as e:
                print(f"⚠️ Could not process image {file}: {e}")
    return documents


# Load DOCX (Word) files
def load_docx(folder_path):
    documents = []
    for file in os.listdir(folder_path):
        if file.lower().endswith(".docx"):
            file_path = os.path.join(folder_path, file)
            try:
                doc = DocxDocument(file_path)
                text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                if text.strip():
                    documents.append(Document(page_content=text, metadata={"source": file}))
                    print(f"📄 Extracted text from DOCX: {file}")
            except Exception as e:
                print(f"⚠️ Could not process DOCX {file}: {e}")
    return documents


# Load all supported documents
def load_documents(folder_path):
    documents = []

    # PDFs
    for file in os.listdir(folder_path):
        if file.lower().endswith(".pdf"):
            file_path = os.path.join(folder_path, file)
            try:
                loader = PyPDFLoader(file_path)
                docs = loader.load()
                if docs:
                    documents.extend(docs)
                else:
                    raise ValueError("No text found in PDF")
            except Exception:
                print(f"🔎 Running OCR on scanned PDF: {file}")
                text = load_scanned_pdf(file_path)
                if text.strip():
                    documents.append(Document(page_content=text, metadata={"source": file}))

    # Images
    documents.extend(load_images(folder_path))

    # DOCX
    documents.extend(load_docx(folder_path))

    return documents


# Split into chunks
def split_docs(documents):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=50
    )
    return splitter.split_documents(documents)


# Create vectorstore
def create_vectorstore(docs):
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    return FAISS.from_documents(docs, embeddings)


# Ask ZhipuAI
def ask_zhipu(question, context):
    client = ZhipuAI(api_key=os.getenv("ZHIPUAI_API_KEY"))
    response = client.chat.completions.create(
        model="glm-4.5-air",
        messages=[
            {"role": "system", "content": "You are a helpful assistant that answers based on provided context."},
            {"role": "user", "content": f"Context:\n{context}\n\nQuestion: {question}"}
        ]
    )
    return response.choices[0].message.content  # ✅ fixed


if __name__ == "__main__":
    folder_path = input("Enter path to your documents folder: ").strip()

    if not os.path.exists(folder_path):
        print(f"❌ The folder '{folder_path}' does not exist. Please try again.")
        exit()

    print("📂 Loading documents...")
    docs = load_documents(folder_path)

    print("✂️ Splitting text into chunks...")
    split_documents = split_docs(docs)

    if not split_documents:
        print("❌ No text found in documents (PDFs, DOCX, Images).")
        exit()

    print("📦 Creating vector store...")
    vectorstore = create_vectorstore(split_documents)
    retriever = vectorstore.as_retriever(search_kwargs={"k": 4})

    print("🤖 Setting up QnA bot with ZhipuAI (GLM-4.5 Air)...")
    print("\n✅ Doc QnA Bot is ready! Type your questions below (or 'exit' to quit).\n")

    while True:
        query = input("You: ")
        if query.lower() in ["exit", "quit"]:
            print("👋 Goodbye!")
            break

        retrieved_docs = retriever.invoke(query)
        context = "\n\n".join([d.page_content for d in retrieved_docs])

        answer = ask_zhipu(query, context)

        print("\nBot:", answer, "\n")

        if retrieved_docs:
            print("📖 Sources:")
            for i, d in enumerate(retrieved_docs, 1):
                print(f"  {i}. {d.metadata.get('source', 'unknown')} (page {d.metadata.get('page', '?')})")
            print()
